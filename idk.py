import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import qn, nsdecls

def create_igcse_notes():
    doc = docx.Document()
    
    # Page Setup
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Palette definition
    COLOR_PRIMARY = RGBColor(0x1B, 0x36, 0x5D)   # Deep Navy
    COLOR_SECONDARY = RGBColor(0x4A, 0x77, 0x7A) # Slate Teal
    COLOR_TEXT = RGBColor(0x22, 0x22, 0x22)      # Charcoal
    COLOR_MUTED = RGBColor(0x55, 0x55, 0x55)     # Muted Grey

    # Configure Styles
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Calibri'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = COLOR_TEXT

    # Helpers
    def add_title(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(12)
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(26)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        return p

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = COLOR_PRIMARY
        
        # Bottom accent border line XML rule
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(r'<w:pBdr xmlns:w="http://openxmlformats.org">'
                         r'<w:bottom w:val="single" w:sz="12" w:space="4" w:color="1B365D"/>'
                         r'</w:pBdr>')
        pPr.append(pBdr)
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(14)
        run.font.bold = True
        run.font.color.rgb = COLOR_SECONDARY
        return p

    def add_bullet(bold_prefix, text_content):
        p = doc.add_paragraph(style='List Bullet')
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(bold_prefix)
        r1.bold = True
        r1.font.color.rgb = COLOR_TEXT
        r2 = p.add_run(text_content)
        r2.font.color.rgb = COLOR_TEXT
        return p

    def add_callout(text):
        table = doc.add_table(rows=1, cols=1)
        table.alignment = docx.enum.table.WD_TABLE_ALIGNMENT.CENTER
        table.autofit = False
        
        cell = table.cell(0, 0)
        cell.width = Inches(6.5)
        
        # Set shading and left border
        tcPr = cell._tc.get_or_add_tcPr()
        tcBorders = parse_xml(r'<w:tcBorders xmlns:w="http://openxmlformats.org">'
                              r'<w:left w:val="single" w:sz="24" w:space="0" w:color="4A777A"/>'
                              r'<w:top w:val="none"/>'
                              r'<w:right w:val="none"/>'
                              r'<w:bottom w:val="none"/>'
                              r'</w:tcBorders>')
        shd = parse_xml(r'<w:shd xmlns:w="http://openxmlformats.org" w:fill="F4F7F7"/>')
        tcPr.append(tcBorders)
        tcPr.append(shd)
        
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.left_indent = Inches(0.1)
        p.paragraph_format.right_indent = Inches(0.1)
        run = p.add_run(text)
        run.font.italic = True
        run.font.color.rgb = COLOR_MUTED
        doc.add_paragraph().paragraph_format.space_after = Pt(6)

    # --- Document Compilation ---
    add_title("IGCSE English Revision Master Guide")
    
    # Section 1
    add_h1("Section 1: English Language (0500) — Composition Frameworks")
    add_h2("Reading Comprehension & Directed Writing")
    add_bullet("Objective: ", "Synthesize context from texts and transform them into a targeted format.")
    add_bullet("Key Criteria: ", "Incorporate 10-12 discrete textual points using completely fresh vocabulary expressions.")
    add_bullet("Structure: ", "Introduction establishing the context, two structured evidence-focused paragraphs, and a definitive conclusion.")
    
    add_h2("Journal Entry")
    add_bullet("Tone: ", "First-person reflective narrative using rich personal pronouns and interior monologues.")
    add_bullet("Structure: ", "Anchor with a date stamp, evaluate immediate emotional shifts, and close with future considerations.")
    
    add_h2("Interview Writing")
    add_bullet("Format: ", "Speaker names must be clearly defined in the left margin without any quotation marks around lines.")
    add_bullet("Structure: ", "Frame with a short italicized contextual setup, followed by a series of structured open-ended questions.")
    
    add_h2("Summary Writing")
    add_bullet("Format: ", "Strictly target a maximum of 120 words using objective, third-person paraphrasing.")
    add_bullet("Constraint: ", "Omit all descriptions, repetitive statements, adjectives, and personal opinions.")
    
    add_h2("Persuasive Speech")
    add_bullet("Rhetorical Toolkit: ", "Incorporate AFOREST structures including tripling, rhetorical questions, and emotional hooks.")
    add_bullet("Structure: ", "Formal opening greeting, clear presentation of the core problem, a viable solution, and a direct call to action.")

    # Section 2
    add_h1("Section 2: English Literature (0475) — Poetry Notes")
    
    add_h2("1. There is a Pleasure in the Pathless Woods (Lord Byron)")
    add_bullet("Themes: ", "The power of the sublime nature, liberation found in isolation, and the permanence of the ocean over human frailty.")
    add_bullet("Literary Devices: ", "Oxymoron ('Society, where none intrudes') and structural Personification ('Roll on, thou deep and dark blue Ocean').")
    add_callout("Analysis: The use of soft p-sound alliterative clusters reinforces the serene, untouched peace of nature.")
    
    add_h2("2. The Way of the World (William Congreve)")
    add_bullet("Themes: ", "Social hypocrisy, shallow societal manners, and the cold economic transactional nature of courtship.")
    add_bullet("Literary Devices: ", "Extended metaphor comparing social interaction to games of chess and sharp satirical irony.")

    add_h2("3. Ulysses (Alfred, Lord Tennyson)")
    add_bullet("Themes: ", "The relentless search for hidden knowledge and the conflict between physical aging and mental willpower.")
    add_bullet("Literary Devices: ", "Metaphorical phrasing ('drink life to the lees') paired with enjambment to mirror an endless journey.")

    add_h2("4. Snake (D.H. Lawrence)")
    add_bullet("Themes: ", "The clash between institutional education and internal nature, alongside the guilt of small-minded human actions.")
    add_bullet("Literary Devices: ", "Onomatopoeia mimicking soft movement, animal similes, and biblical/regal allusions.")

    # Section 3
    add_h1("Section 3: English Literature — Prose Notes")
    
    add_h2("1. The Open Window (Saki)")
    add_bullet("Themes: ", "The chaotic power of dynamic storytelling and how rigid politeness leaves individuals vulnerable.")
    add_bullet("Character Traits: ", "Vera is calculative and highly imaginative. Framton Nuttel is anxious and easily manipulated.")
    add_bullet("Literary Devices: ", "Framed inner narrative structure and situational irony that subverts structural expectations.")

    add_h2("2. The Cactus (O. Henry)")
    add_bullet("Themes: ", "The dangers of ego and vanity, combined with major breakdowns in interpersonal communication.")
    add_bullet("Character Traits: ", "Trysdale is narcissistic, self-absorbed, and blind to clear situational clues.")
    add_bullet("Literary Devices: ", "Dramatic irony, shifting twist ending, and the cactus serving as a central symbol of missed affection.")

    add_h2("3. A Widow")
    add_bullet("Themes: ", "The painful isolation of grief and finding resilience against cold societal expectations.")
    add_bullet("Character Traits: ", "The Widow is quiet, carrying deep internal grief, yet independent in her choices.")
    add_bullet("Literary Devices: ", "Pathetic fallacy linking landscapes to internal feelings, contrasted with social dialogue blocks.")

    # Section 4
    add_h1("Section 4: English Literature — Drama (Julius Caesar)")
    add_bullet("Act 3 (The Crisis): ", "Caesar's death in the Senate followed by Antony's emotionally manipulative funeral speech.")
    add_bullet("Act 4 (The Decline): ", "The cold planning of the new Triumvirate and the bitter argument between Brutus and Cassius.")
    add_bullet("Act 5 (The Catastrophe): ", "The final defeat at Philippi, leading to suicides that preserve tragic honor.")
    
    add_h2("Core Structural Themes")
    add_bullet("Key Elements: ", "The dangerous power of skilled political rhetoric, fate versus free will, and the failure of idealism when facing pragmatism.")
    
    add_h2("Character Analysis")
    add_bullet("Brutus: ", "Stoic, driven by rigid principles, patriotic, but critically blind to human manipulation.")
    add_bullet("Mark Antony: ", "Machiavellian strategist, fiercely loyal, and incredibly skilled at reading crowd psychology.")
    
    add_h2("Key Literary Mechanisms")
    add_bullet("Devices: ", "The structural switch between blank verse and prose, calculated verbal irony, and deep soliloquy structures.")
    add_callout("Analysis: Antony's repetitive shift of the word 'honorable' acts as structural antiphrasis, turning the crowd into an angry mob.")

    doc.save("IGCSE_English_Comprehensive_Study_Notes.docx")

if __name__ == "__main__":
    create_igcse_notes()
